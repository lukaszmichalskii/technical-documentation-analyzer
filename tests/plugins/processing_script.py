from __future__ import annotations

import os
import sys
import pathlib
import types
from abc import abstractmethod
import abc
import re
import typing

import PyPDF2
import docx


class NotSupportedDocumentFormat(Exception):
    pass


def _read_all(
        fd, buffer_read_fn: typing.Callable[[typing.IO[str | bytes]], str]
) -> str:
    """
    Read from buffer stream str or bytes and clean up decoded content
    Args:
        fd: file descriptor
        buffer_read_fn: function to read buffer stream
    Returns:
        text without escape characters
    """
    return remove_escape_chars(buffer_read_fn(fd))


class TextProvider:
    """
    Class for handling single input - single output by reading textual content from files.
    """

    def __init__(self) -> None:
        self.supported_formats = {".docx", ".pdf"}
        self.pdf_extractor = PDFDecoder()
        self.docx_extractor = DocxDecoder()

    def supported_format(self, extension) -> bool:
        return extension in self.supported_formats

    def get_file_content(self, filepath: pathlib.Path) -> str:
        """
        Read whole single .pdf, .docx file to program memory
        Args:
            filepath: path to file
        Returns:
            decoded and cleaned text from provided file.
        """
        if not self.supported_format(filepath.suffix):
            raise NotSupportedDocumentFormat(
                f"Document format {filepath.suffix} is not supported."
            )
        if filepath.suffix == ".pdf":
            with open(filepath, "rb") as pdf_:
                return _read_all(pdf_, self.pdf_extractor.read_all)
        elif filepath.suffix == ".docx":
            with open(filepath, "rb") as docx_:
                return _read_all(docx_, self.docx_extractor.read_all)

    def get_file_chunk(
            self, filepath: pathlib.Path
    ) -> typing.Generator[str, None, None]:
        """
        Read single .pdf, .docx, and standard .txt file content chunks
        Args:
            filepath: path to file
        Returns:
            decoded and cleaned text chunks from provided file.
        """
        if not self.supported_format(filepath.suffix):
            raise NotSupportedDocumentFormat(
                f"Document format {filepath.suffix} is not supported."
            )
        if filepath.suffix == ".pdf":
            with open(filepath, "rb") as pdf_:
                while True:
                    try:
                        yield remove_escape_chars(next(self.pdf_extractor.read(pdf_)))
                    except StopIteration:
                        self.pdf_extractor.reset()
                        break
        elif filepath.suffix == ".docx":
            with open(filepath, "rb") as docx_:
                while True:
                    try:
                        yield remove_escape_chars(next(self.docx_extractor.read(docx_)))
                    except StopIteration:
                        self.docx_extractor.reset()
                        break


def remove_escape_chars(string_: str) -> str:
    """
    Strip string with escape characters e.g. tabs, newlines, terminated chars
    Args:
        string_: text to clean up
    Returns:
        parsed text without escape characters
    """
    escape_chars = r"\s+"
    return re.sub(escape_chars, " ", string_).strip()


class Decoder(abc.ABC):
    """
    Base class for files decoding
    """

    def __init__(self) -> None:
        self.head = -1

    @abstractmethod
    def read(self, fd: typing.IO[str | bytes]) -> typing.Generator[str, None, None]:
        pass

    @abstractmethod
    def reset(self) -> None:
        pass

    @staticmethod
    def read_all(fd: typing.IO[str | bytes]) -> str:
        pass


class PDFDecoder(Decoder):
    """
    Decoder class providing method for reading .pdf based content only.
    """

    def __init__(self) -> None:
        super().__init__()

    def read(self, fd: typing.IO[str | bytes]) -> typing.Generator[str, None, None]:
        """
        Decode and extract PDF file content page by page
        Args:
            fd: PDF file descriptor
        Returns:
            decoded page content
        """
        pdf_reader = PyPDF2.PdfReader(fd)
        self.head += 1
        while self.head < len(pdf_reader.pages):
            yield pdf_reader.pages[self.head].extract_text()

    @staticmethod
    def read_all(fd: typing.IO[str | bytes]) -> str:
        text = ""
        pdf_reader = PyPDF2.PdfReader(fd)
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text

    def reset(self) -> None:
        self.head = -1


class DocxDecoder(Decoder):
    """
    Decoder class providing method for reading .docx based content only.
    By now class does not take into consideration images, tables and other MS Word structures.
    """

    def __init__(self) -> None:
        super().__init__()

    def read(self, fd: typing.IO[str | bytes]) -> typing.Generator[str, None, None]:
        """
        Decode and extract .docx file content paragraph by paragraph.
        Args:
            fd: path to .docx encoded file
        Returns:
            decoded paragraph content
        """
        docx_ = docx.Document(fd)
        self.head += 1
        while self.head < len(docx_.paragraphs):
            yield docx_.paragraphs[self.head].text

    @staticmethod
    def read_all(fd: typing.IO[str | bytes]) -> str:
        text = ""
        docx_ = docx.Document(fd)
        for paragraph in docx_.paragraphs:
            text += paragraph.text
        return text

    def reset(self) -> None:
        self.head = -1


def save_parsed_text(
        destination: pathlib.Path,
        parsed_text: str | typing.Generator[str, None, None],
) -> None:
    with open(destination, "w") as fd:
        if isinstance(parsed_text, types.GeneratorType):
            for text in parsed_text:
                fd.write(text)
        else:
            fd.write(parsed_text)


if __name__ == '__main__':
    args = sys.argv[1:]

    file_size_limit = int(os.environ.get("IN_MEMORY_FILE_SIZE", 1024 * 1024))
    text_provider = TextProvider()

    file = pathlib.Path(args[0])

    decoded_text = text_provider.get_file_content(file)

    save_parsed_text(pathlib.Path(args[1]), decoded_text)
