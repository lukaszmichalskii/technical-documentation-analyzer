from __future__ import annotations

import abc
import re
import typing
from abc import abstractmethod

import PyPDF2
import docx


def read_file(
    file_obj: typing.IO[str | bytes], chunk_size: int = 1024 * 32
) -> typing.Generator[str, None, None]:
    """
    Read file content to memory buffer
    Args:
        file_obj: file descriptor
        chunk_size: memory buffer size, default is 32KiB
    Returns:
         buffer filled with file_obj content
    """
    while True:
        text = file_obj.read(chunk_size)
        if not text:
            break
        yield text


def remove_escape_chars(string_: str) -> str:
    """
    Strip string with escape characters e.g. tabs, newlines, terminated chars
    Args:
        string_: text to clean up
    Returns:
        parsed text without escape characters
    """
    escape_chars = r"\s+"
    return re.sub(escape_chars, " ", string_).strip()


class Decoder(abc.ABC):
    """
    Base class for files decoding
    """

    def __init__(self) -> None:
        self.head = -1

    @abstractmethod
    def read(self, fd: typing.IO[str | bytes]) -> typing.Generator[str, None, None]:
        pass

    @abstractmethod
    def reset(self) -> None:
        pass

    @staticmethod
    def read_all(fd: typing.IO[str | bytes]) -> str:
        pass


class PDFDecoder(Decoder):
    """
    Decoder class providing method for reading .pdf based content only.
    """

    def __init__(self) -> None:
        super().__init__()

    def read(self, fd: typing.IO[str | bytes]) -> typing.Generator[str, None, None]:
        """
        Decode and extract PDF file content page by page
        Args:
            fd: PDF file descriptor
        Returns:
            decoded page content
        """
        pdf_reader = PyPDF2.PdfReader(fd)
        self.head += 1
        while self.head < len(pdf_reader.pages):
            yield pdf_reader.pages[self.head].extract_text()

    @staticmethod
    def read_all(fd: typing.IO[str | bytes]) -> str:
        text = ""
        pdf_reader = PyPDF2.PdfReader(fd)
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text

    def reset(self) -> None:
        self.head = -1


class DocxDecoder(Decoder):
    """
    Decoder class providing method for reading .docx based content only.
    By now class does not take into consideration images, tables and other MS Word structures.
    """

    def __init__(self) -> None:
        super().__init__()

    def read(
        self, fd: typing.IO[str | bytes]
    ) -> typing.Generator[str, None, None]:
        """
        Decode and extract .docx file content paragraph by paragraph.
        Args:
            fd: path to .docx encoded file
        Returns:
            decoded paragraph content
        """
        docx_ = docx.Document(fd)
        self.head += 1
        while self.head < len(docx_.paragraphs):
            yield docx_.paragraphs[self.head].text

    @staticmethod
    def read_all(fd: typing.IO[str | bytes]) -> str:
        text = ""
        docx_ = docx.Document(fd)
        for paragraph in docx_.paragraphs:
            text += paragraph.text
        return text

    def reset(self) -> None:
        self.head = -1
